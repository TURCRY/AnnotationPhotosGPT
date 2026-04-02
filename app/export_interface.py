import streamlit as st
import pandas as pd
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INFOS_PROJET_PATH = "data/infos_projet.json"
ANNOTATION_CSV_PATH = "data/annotations.csv"

def load_infos_projet():
    with open(INFOS_PROJET_PATH, "r", encoding="utf-8") as f:
        return json.load(f)

def export_interface():
    st.title("📤 Export des résultats")

    infos = load_infos_projet()

    if not Path(ANNOTATION_CSV_PATH).exists():
        st.warning("Aucune annotation trouvée à exporter.")
        return

    df = pd.read_csv(ANNOTATION_CSV_PATH, sep=";")

    st.markdown("### 📑 Aperçu des annotations")
    st.dataframe(df)

    nom_export = st.text_input("Nom de fichier Word exporté (sans extension)", value="export_annotation")

    if st.button("📝 Exporter en Word avec images et légendes"):
        doc = Document()
        doc.add_heading("Annotations de photos", level=1)

        for i, row in df.iterrows():
            table = doc.add_table(rows=1, cols=2)
            table.allow_autofit = True

            cell_comment = table.cell(0, 0)
            cell_comment.text = str(row['commentaire'])
            cell_comment.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell_comment.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.style.font.size = Pt(11)

            cell_photo = table.cell(0, 1)
            cell_photo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell_photo.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            try:
                run = para.add_run()
                run.add_picture(row['chemin_photo_reduite'], width=Inches(2.5))

                # Ajouter une légende après l’image
                legend_paragraph = cell_photo.add_paragraph()
                legend_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                legend_paragraph.style = 'Caption'
                legend_paragraph.add_run(f"Cliché {i + 1} – {row['libelle']}")
            except Exception as e:
                para.add_run("Image non trouvée")

            doc.add_paragraph("")

        output_path = Path("dist")
        output_path.mkdir(exist_ok=True)
        file_path = output_path / f"{nom_export}.docx"
        doc.save(file_path)

        st.success(f"Fichier Word exporté vers : {file_path}")