from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
import pandas as pd
import os, glob, json
from PIL import Image
from datetime import datetime
from pathlib import Path
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import unicodedata

# --- Helpers ----------------------------------------------------------------
def load_photos(path: Path) -> pd.DataFrame:
    if path.suffix.lower() == ".xlsx":
        return pd.read_excel(path, engine="openpyxl")
    try:
        return pd.read_csv(path, sep=";", encoding="utf-8-sig")
    except UnicodeDecodeError:
        return pd.read_csv(path, sep=";", encoding="latin-1")


def coerce_bool_series(s: pd.Series) -> pd.Series:
    return s.astype(str).str.strip().str.lower().isin(["true","1","yes","y","oui"])

def find_latest_annotations(base_dir: Path) -> Path:
    cands = sorted(base_dir.glob("*_GTP_*.csv"), key=os.path.getmtime, reverse=True)
    if not cands:
        raise FileNotFoundError("Aucune annotation *_GTP_*.csv trouvée.")
    return cands[0]

def safe_text(val: object) -> str:
    """Convertit proprement une valeur (NaN, float, None...) en texte."""
    try:
        import pandas as _pd
        if val is None:
            return ""
        if isinstance(val, float) and _pd.isna(val):
            return ""
    except Exception:
        # si pandas pas dispo ici, on se contente de gérer None
        if val is None:
            return ""
    return str(val)


def normalize(s: str) -> str:
    return unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode().lower().strip()

def iter_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def find_marker_paragraph(doc: Document, marker: str) -> Paragraph:
    for p in iter_paragraphs(doc):
        if marker in p.text:
            return p
    raise ValueError(f"Marqueur introuvable dans le modèle : {marker}")

def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    if text:
        p.add_run(text)
    return p

def add_field_run(paragraph: Paragraph, instr: str, placeholder: str = ""):
    """
    Ajoute un champ Word (Field) dans un paragraphe.
    instr: ex. 'SEQ Cliché \\* ARABIC' ou 'TOC \\h \\z \\c "Cliché"'
    placeholder: texte affiché tant que Word n'a pas mis à jour les champs.
    """
    run = paragraph.add_run()

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = instr

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    # Texte provisoire affiché avant mise à jour des champs dans Word
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = placeholder

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    run._r.extend([fldChar1, instrText, fldChar2, fldChar3, fldChar4])
    return run


def add_caption_cliche(doc: Document, libelle: str, style_name: str = "Légende"):
    """
    Ajoute une légende Word 'Cliché X - ...' basée sur un champ SEQ, conforme à Word.
    La table des clichés pourra être générée via TOC \\c "Cliché".
    """
    p = doc.add_paragraph()

    # (Optionnel mais recommandé) appliquer le style de légende du modèle
    # => uniquement pour la mise en forme, pas pour l’indexation
    try:
        p.style = doc.styles[style_name]
    except Exception:
        pass

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # ou CENTER si votre modèle le veut

    # "Cliché " + {SEQ Cliché} + " - " + libellé
    p.add_run("Cliché ")
    add_field_run(p, r'SEQ Cliché \* ARABIC', placeholder="1")
    p.add_run(" - " + libelle)

    return p



# --- Entrées ----------------------------------------------------------------
REPORT_MODE = os.environ.get("REPORT_MODE", "UI").upper().strip()  # UI / GTP
ONLY_RETENUE = os.environ.get("REPORT_ONLY_RETENUE", "1").strip().lower() in ("1", "true", "yes", "oui")

# On part de la structure du projet : scripts/ .. / data/
project_root = Path(__file__).resolve().parents[1]

# 1) infos_projet.json
infos_path = project_root / "data" / "infos_projet.json"
if not infos_path.exists():
    raise FileNotFoundError(f"infos_projet.json introuvable : {infos_path}")
with open(infos_path, encoding="utf-8") as f:
    infos = json.load(f)

# 2) CSV photos UI
photos_path_str = infos.get("fichier_photos", "")
if not photos_path_str:
    raise FileNotFoundError("fichier_photos manquant dans infos_projet.json")
photos_path = Path(photos_path_str)
if not photos_path.exists():
    raise FileNotFoundError(f"Fichier photos introuvable : {photos_path}")
photos_df = load_photos(photos_path)

base_dir = photos_path.parent

# 3) Charger GTP si nécessaire (ou si on veut les textes GTP)

try:
    annotations_path = find_latest_annotations(base_dir)
    try:
        annotations_df = pd.read_csv(annotations_path, sep=";", encoding="utf-8-sig")
    except UnicodeDecodeError:
        annotations_df = pd.read_csv(annotations_path, sep=";", encoding="latin-1")
except FileNotFoundError:
    annotations_df = None
    annotations_path = None


# 4) Charger batch si dispo
batch_df = None
photos_batch_path = infos.get("fichier_photos_batch", "")
if photos_batch_path and Path(photos_batch_path).exists():
    batch_df = pd.read_csv(photos_batch_path, sep=";", encoding="utf-8-sig")

# 5) Base = UI (source de vérité pour la liste des photos)
df = photos_df.copy()

# normaliser retenue
if "retenue" in df.columns:
    df["retenue"] = coerce_bool_series(df["retenue"])
else:
    df["retenue"] = True

# --- Filtrage par mode -------------------------------------------------------
if REPORT_MODE == "GTP":
    if annotations_df is None or annotations_df.empty:
        raise RuntimeError("Mode GTP : aucun *_GTP_*.csv chargé / vide.")

    if "nom_fichier_image" not in annotations_df.columns:
        raise RuntimeError("Mode GTP : colonne 'nom_fichier_image' absente du *_GTP_*.csv")

    ann = annotations_df.copy()
    ann["nom_fichier_image"] = ann["nom_fichier_image"].astype(str).str.strip()

    if "annotation_validee" in ann.columns:
        valid = pd.to_numeric(ann["annotation_validee"], errors="coerce").fillna(0).astype(int) == 1
        gtp_names = set(ann.loc[valid, "nom_fichier_image"])
    else:
        # fallback : on prend toutes les lignes GTP si pas de colonne
        gtp_names = set(ann["nom_fichier_image"])

    df["nom_fichier_image"] = df["nom_fichier_image"].astype(str).str.strip()
    df = df[df["nom_fichier_image"].isin(gtp_names)].copy()

# Filtre retenue (dans les deux modes, si demandé)
if ONLY_RETENUE and "retenue" in df.columns:
    df = df[df["retenue"]].copy()

# --- Merge GTP (si mode GTP) ------------------------------------------------
if annotations_df is not None and not annotations_df.empty:
    wanted_gtp = ["nom_fichier_image", "libelle", "commentaire", "retenue"]
    cols_gtp = [c for c in wanted_gtp if c in annotations_df.columns]
    if cols_gtp:
        ann2 = annotations_df[cols_gtp].copy()
        ann2["nom_fichier_image"] = ann2["nom_fichier_image"].astype(str).str.strip()
        df["nom_fichier_image"] = df["nom_fichier_image"].astype(str).str.strip()
        df = df.merge(ann2, on="nom_fichier_image", how="left", suffixes=("", "_gtp"))

# --- Merge batch -------------------------------------------------------------
if batch_df is not None and not batch_df.empty:
    if "photo_rel_native" in df.columns and "photo_rel_native" in batch_df.columns:
        b = batch_df.copy()
        b["photo_rel_native"] = b["photo_rel_native"].astype(str).str.strip()
        df["photo_rel_native"] = df["photo_rel_native"].astype(str).str.strip()
        df = df.merge(
            b[["photo_rel_native","libelle_propose_batch","commentaire_propose_batch","batch_status","batch_ts"]],
            on="photo_rel_native",
            how="left",
            suffixes=("", "_batch"),
        )

# --- Champs finaux : priorité GTP > UI > Batch -------------------------------
def _non_empty(series: pd.Series) -> pd.Series:
    return series.notna() & (series.astype(str).str.strip() != "")

df["libelle_final"] = ""
df["commentaire_final"] = ""

# GTP
if "libelle" in df.columns:
    df.loc[_non_empty(df["libelle"]), "libelle_final"] = df["libelle"].astype(str)
if "commentaire" in df.columns:
    df.loc[_non_empty(df["commentaire"]), "commentaire_final"] = df["commentaire"].astype(str)

# UI
if "libelle_propose_ui" in df.columns:
    m = (df["libelle_final"].astype(str).str.strip() == "") & _non_empty(df["libelle_propose_ui"])
    df.loc[m, "libelle_final"] = df["libelle_propose_ui"].astype(str)
if "commentaire_propose_ui" in df.columns:
    m = (df["commentaire_final"].astype(str).str.strip() == "") & _non_empty(df["commentaire_propose_ui"])
    df.loc[m, "commentaire_final"] = df["commentaire_propose_ui"].astype(str)

# Batch
if "libelle_propose_batch" in df.columns:
    m = (df["libelle_final"].astype(str).str.strip() == "") & _non_empty(df["libelle_propose_batch"])
    df.loc[m, "libelle_final"] = df["libelle_propose_batch"].astype(str)
if "commentaire_propose_batch" in df.columns:
    m = (df["commentaire_final"].astype(str).str.strip() == "") & _non_empty(df["commentaire_propose_batch"])
    df.loc[m, "commentaire_final"] = df["commentaire_propose_batch"].astype(str)

# provenance
df["source_texte"] = "VIDE"
has_batch = pd.Series(False, index=df.index)
if "libelle_propose_batch" in df.columns:
    has_batch |= _non_empty(df["libelle_propose_batch"])
if "commentaire_propose_batch" in df.columns:
    has_batch |= _non_empty(df["commentaire_propose_batch"])



has_ui = pd.Series(False, index=df.index)
if "libelle_propose_ui" in df.columns:
    has_ui |= _non_empty(df["libelle_propose_ui"])
if "commentaire_propose_ui" in df.columns:
    has_ui |= _non_empty(df["commentaire_propose_ui"])

has_gtp = pd.Series(False, index=df.index)
if "libelle" in df.columns:
    has_gtp |= _non_empty(df["libelle"])
if "commentaire" in df.columns:
    has_gtp |= _non_empty(df["commentaire"])

df.loc[has_batch, "source_texte"] = "BATCH"
df.loc[has_ui,    "source_texte"] = "UI"
df.loc[has_gtp,   "source_texte"] = "GTP"

contexte = {}
ctx_path = infos.get("fichier_contexte_general", "contexte_general.json")
ctx_path = Path(ctx_path)
if not ctx_path.is_absolute():
    ctx_path = project_root / ctx_path
if ctx_path.exists():
    with open(ctx_path, encoding="utf-8") as f:
        contexte = json.load(f)


# --- Document Word -----------------------------------------------------------
TEMPLATE_NAME = "Modele word rapport ver 15 05 2025.docx"
template_path = project_root / "data" / TEMPLATE_NAME
if not template_path.exists():
    raise FileNotFoundError(f"Modèle Word introuvable : {template_path}")

doc = Document(str(template_path))

if "horodatage_photo" in df.columns:
    df = df.sort_values("horodatage_photo", kind="stable")
elif "photo_rel_native" in df.columns:
    df = df.sort_values("photo_rel_native", kind="stable")

# En-tête d'informations
doc.add_paragraph("🔎 Informations générales").bold = True
doc.add_paragraph(f"Annotations : {annotations_path.name if annotations_path else '— (source UI)'}")
doc.add_paragraph(f"Photos : {Path(photos_path).name}")
doc.add_paragraph(f"Généré le : {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
doc.add_paragraph(f"Utilisateur : {infos.get('user', 'N/A')}")
doc.add_paragraph(f"Modèle LLM : {infos.get('model', 'N/A')}")
doc.add_paragraph(f"Mission : {infos.get('mission', 'Non renseignée')}")
doc.add_paragraph(f"Prompt système : {contexte.get('system', 'N/A')}")
doc.add_paragraph("")

# Table des clichés (TOC des légendes)
MARKER = "[[RAPPORT_PHOTOS]]"
p0 = find_marker_paragraph(doc, MARKER)
p0.text = p0.text.replace(MARKER, "").strip()  # nettoie le marqueur dans le modèle

# 1) Table des clichés (TOC des légendes)
p_title = insert_paragraph_after(p0, "Table des clichés :")
p_title.runs[0].bold = True

p_field = insert_paragraph_after(p_title)
run = p_field.add_run()
fldChar1 = docx.oxml.OxmlElement('w:fldChar'); fldChar1.set(docx.oxml.ns.qn('w:fldCharType'), 'begin')
instrText = docx.oxml.OxmlElement('w:instrText'); instrText.set(docx.oxml.ns.qn('xml:space'), 'preserve')
instrText.text = 'TOC \\h \\z \\c "Cliché"'
fldChar2 = docx.oxml.OxmlElement('w:fldChar'); fldChar2.set(docx.oxml.ns.qn('w:fldCharType'), 'separate')
fldChar3 = docx.oxml.OxmlElement('w:t'); fldChar3.text = "Table des clichés (mettre à jour les champs dans Word)"
fldChar4 = docx.oxml.OxmlElement('w:fldChar'); fldChar4.set(docx.oxml.ns.qn('w:fldCharType'), 'end')
run._r.extend([fldChar1, instrText, fldChar2, fldChar3, fldChar4])

p_cursor = insert_paragraph_after(p_field)  # curseur d’insertion pour le corps



# Corps : une table (commentaire + image) par cliché
for num, row in enumerate(df.itertuples(index=False), start=1):
    commentaire = safe_text(getattr(row, "commentaire_final", ""))
    libelle = safe_text(getattr(row, "libelle_final", ""))
    orientation = str(getattr(row, "orientation_photo", "0") or "0")
    base_dir_img = safe_text(getattr(row, "chemin_photo_reduite", "")) or safe_text(getattr(row, "chemin_photo_native", ""))
    nom_fichier = safe_text(getattr(row, "nom_fichier_image", ""))


    photo_path = os.path.normpath(os.path.join(base_dir_img, nom_fichier))

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3.5)
    table.columns[1].width = Inches(3.5)

    # supprime bordures visuelles
    for cell in table.row_cells(0):
        for border in cell._element.xpath(".//w:tcBorders"):
            cell._element.remove(border)

    # Colonne commentaire
    cell_comment = table.cell(0, 0)
    p = cell_comment.paragraphs[0]
    p.text = commentaire
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tcPr = cell_comment._tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign'); vAlign.set(qn('w:val'), "center")
    tcPr.append(vAlign)

    # Colonne photo
    cell_photo = table.cell(0, 1)
    paragraph = cell_photo.paragraphs[0]
    run = paragraph.add_run()
    if os.path.exists(photo_path):
        try:
            img = Image.open(photo_path)
            if orientation in ["90", "270", "180"]:
                img = img.rotate(-int(orientation), expand=True)
            temp_path = f"temp_{num}.jpg"
            img.save(temp_path)
            if orientation in ["90", "270"]:
                run.add_picture(temp_path, height=Inches(2.36))  # ≈6 cm
            else:
                run.add_picture(temp_path, width=Inches(3.2))
            os.remove(temp_path)
        except Exception:
            paragraph.add_run("[Erreur image]").bold = True
    else:
        paragraph.add_run("[Image introuvable]").bold = True

    # Légende (avec timecode si dispo)
    # Ligne vide pour lisibilité (entre tableau photo/commentaire et légende)
    doc.add_paragraph("")

    # Légende Word native (Caption)
    add_caption_cliche(doc, libelle, style_name="Légende")
    source = safe_text(getattr(row, "source_texte", ""))
    if source and source != "VIDE":
        p = doc.add_paragraph()
        p.add_run(f"Source : {source}")
        p.style = doc.styles["Normal"]

    # Ligne vide après
    doc.add_paragraph("")

# Sortie — nom normalisé
id_affaire   = str(infos.get("id_affaire") or "").strip()
id_captation = str(infos.get("id_captation") or "").strip()

# horodatage de génération (jour_heure)
ts = datetime.now().strftime("%Y-%m-%d_%H-%M")

report_name = f"annotation_photos_{id_affaire}_{id_captation}_V_{ts}.docx"

# même dossier que le CSV GTP (comportement actuel)

output_path = base_dir / report_name


doc.save(output_path)
print(f"✅ Rapport généré : {output_path}")


